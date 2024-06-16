
from django.http import HttpResponse, HttpResponseForbidden
from docx import Document
import pandas as pd
from django.conf import settings
from copy import *

'''@usuario_de_setor_especifico('Tecnologia da Informação')''' # Decorador para controle de acesso por base em dados do banco.

'''@login_required(login_url="/GIT/login/")''' # Decorador para controle de acesso para usuarios logados.


# Criando no django, um decorador vinculado a uma informação especifica do Banco de Dados.
def usuario_de_setor_especifico(setor):
    def decorador(view_func):
        @wraps(view_func)
        def visualizacao_envolvida(request, *args, **kwargs):
            usuario = Dados.objects.get(username=request.user.username)
            if usuario.setor == setor:
                return view_func(request, *args, **kwargs)
            else:
                e ='A Página foi bloquada para seu usuario.'
                return redirect(error_handler, e) 
        return visualizacao_envolvida
    return decorador

# Ativar e Desativar contas.
def alterar_status(request, username):
        usuario = User.objects.get(username = username)
        atual = usuario.is_active
        if atual == 1:
            usuario.is_active = 0
            logger.warning(f'{request.user} desativou a conta do usuario {username}')
        else:
            usuario.is_active = 1
            logger.warning(f'{request.user} ativou a conta do usuario {username}')
        usuario.save()

        # Redirecione ou retorne uma resposta de sucesso
        return redirect(listagem_de_usuarios)

# Redefinir Senha para uma padrão.
def redefinir_senha(request, username):
        usuario = Dados.objects.get(username = username)
        logger.warning(f'{request.user} redefiniu a senha da conta do usuario {username} para a senha padrão!')

        # Obtenha a nova senha do formulário
        nova_senha = 'Senha_Padrão'

        # Altere a senha do usuário
        usuario.set_password(nova_senha)
        usuario.save()

        # Redirecione ou retorne uma resposta de sucesso
        return redirect(listagem_de_usuarios)

# Alterando Senha via Inputs.
def alterar_senha(request):
    if request.method == 'POST':
        # Identifique o usuário cuja senha você deseja alterar
        usuario = Dados.objects.get(username = request.user.username)

        # Obtenha a nova senha do formulário
        nova_senha = request.POST.get('password')
        if nova_senha == 'Senha_Padrão':
            return render(request, 'alterar_senha.html', {'erro_code':'Não use a senha padrão!'})
        else:
            # Altere a senha do usuário
            usuario.set_password(nova_senha)
            usuario.save()
            # Redirecione ou retorne uma resposta de sucesso
            return redirect(consultar)

# Listar todos os elementos de uma tabela do banco.
def listagem_de_usuarios(request):
        users = User.objects.all()
        data = Dados.objects.all()
        return render(request, 'listar.html', {'users': users, 'data': data})

# Excluir elementos de uma tabela do banco.
def deletar_usuario(request, pk):
    try:
        # Busca o usuário pelo ID
        user = User.objects.get(id=pk)
        dados = Dados.objects.filter(username=user.username).first()
        # Deleta o usuário
        user.delete()
        if dados:
            dados.delete()
        logger.warning(f'O usuário {request.user.username}, deletou a conta de {user.username} com sucesso!')
    except User.DoesNotExist:
        logger.warning(f'Usuário com o ID {pk} não encontrado.')
    
    return redirect(listagem_de_usuarios)

# Editar dados de um usuario.
def editar_usuarios(request , pk):
    usuario = get_object_or_404(User, pk=pk)

    if request.method == "POST":
        cargo = request.POST.get('cargo')
        setor = request.POST.get('setor')
        empresa = request.POST.get('empresa')
        pronome = "Senhor(a)"

        usuario.last_name = cargo
        usuario.email = f"{usuario.username}@github.org.br"
        usuario.save()

        dados_usuario, created = Dados.objects.get_or_create(username=usuario.username)
        dados_usuario.nome_completo = request.user.first_name
        dados_usuario.cargo = cargo
        dados_usuario.setor = setor
        dados_usuario.empresa = empresa
        dados_usuario.email = usuario.email
        dados_usuario.pronome = pronome
        dados_usuario.cpf = request.user.last_name
        dados_usuario.save()

        logger.error(f"O usuário {request.user}, alterou o setor do usuário {dados_usuario.username} para {dados_usuario.setor}.")
        logger.error(f"O usuário {request.user}, alterou o cargo do usuário {dados_usuario.username} para {dados_usuario.cargo}.")
        logger.error(f"O usuário {request.user}, alterou o local de trabalho do usuário {dados_usuario.username} para {dados_usuario.empresa}.")

        return redirect("listagem_usuarios")
        
    return render(request, 'editar_usuarios.html', {
        'usuario': usuario,
    })

# Deletar pastas e arquivos do servidor.
def deletar_pasta_e_arquivos(pasta):
    # Verifique se a pasta existe
    if os.path.exists(pasta):
        # Exclua todos os arquivos dentro da pasta
        for filename in os.listdir(pasta):
            file_path = os.path.join(pasta, filename)
            try:
                # Se o caminho for um diretório, exclua-o recursivamente
                if os.path.isdir(file_path):
                    shutil.rmtree(file_path)
                # Caso contrário, exclua o arquivo
                else:
                    os.remove(file_path)
            except Exception as e:
                url = reverse('error_handler', kwargs={'exception': str(e)})
                return redirect(url)

        # Exclua a própria pasta
        try:
            shutil.rmtree(pasta)
        except Exception as e:
            url = reverse('error_handler', kwargs={'exception': str(e)})
            return redirect(url)
        
# Deletar um arquivo especifico.
def deletar_arquivo(caminho):
    if os.path.exists(caminho):
        os.remove(caminho)

# Função de Login via LDAP
def login(request):
    if request.user.is_authenticated:
        return redirect('base')

    elif request.method == "GET":
        return render(request, 'login.html', {'erro_code': None})
    
    else:
        username = request.POST.get('username')
        password = request.POST.get('password')

        try:
            # Realiza a autenticação no LDAP utilizando os dados informados
            usuario = ldap_authenticate(username, password)
            if usuario:
                cpf = usuario[1]
                usuario = usuario[0]
                username = username.lower()
                # Verifica se o usuário já existe no banco de dados
                b_usuario = User.objects.filter(username=username).first()
                if b_usuario:
                    if not b_usuario.is_active:
                        return render(request, 'login.html', {'erro_code':  "Sua conta está inativa. Abra chamado para TI."})
                    
                else:
                    logger.warning("Cadastro de um novo Usuario")
                    b_usuario = User.objects.create_user(username=username, password=password, first_name=cpf, is_active=0)
                    b_usuario.save()
                    if not b_usuario.is_active:
                        return render(request, 'login.html', {'erro_code':  "Sua conta está inativa. Entre em contato com o administrador."})
                    
                user = authenticate(username = username, password = password)
                login(request, user)
                return redirect('base')
            else:
                return render(request, 'login.html', {'erro_code' : "Usuario ou senha incorretos!"})
        except:
            user = authenticate(username=username, password=password)
            logger.warning("Falta o CPF no LDAP!")
            return render(request, 'login.html', {'erro_code' : "Falta o CPF no LDAP!"})

# Criando um Documento DOCX
def criar_oficio(resultado, quantidade):
    documento_modelo = 'Caminho_Do_Documento_Modelo'

    for i in range(quantidade):
            documento = Document(documento_modelo)
            novo_documento = f'Caminho_do_novo_documento/word{i}.docx'
            for tabela in documento.tables:
                # Verificar se a tabela possui pelo menos uma linha e duas colunas
                if len(tabela.rows) > 0 and len(tabela.columns) >= 2:
                    # Chamar a função para preencher a segunda coluna com os valores do dicionário
                    preencher_tabela_com_dicionario(tabela, resultado[i])
            documento.save(novo_documento)
        
# Preenchendo dados em um documento DOCX que possio uma tabela, utilizando dicionarios e for's.
def preencher_tabela_com_dicionario(tabela, dicionario):
    # Iterar sobre as linhas da tabela
    for i, linha in enumerate(tabela.rows):
        # Verificar se a linha possui pelo menos duas células
        if len(linha.cells) >= 2:
            # Obter a segunda célula (segunda coluna)
            segunda_celula = linha.cells[1]
            # Verificar se ainda há itens no dicionário
            if i < len(dicionario):
                # Obter o valor correspondente do dicionário
                valor = list(dicionario.values())[i]
                # Preencher a segunda célula com o valor do dicionário
                segunda_celula.text = str(valor)

# Procurando palavras chaves em uma planilha e retornando tudo da mesma linha.
def puxar_dados(cpf=00000000000):

    arquivo =  'Caminho_onde_deve_procurar_o_cpf_em_excel'
    df = pd.read_excel(arquivo)

    df['Nome da Coluna CPF'] = df['Nome da Coluna CPF'].astype(str).str.strip()

    # Use a função any para verificar se algum valor na coluna 'Nome da Coluna' é igual ao CPF desejado
    if (df['Nome da Coluna CPF'] == cpf).any():
        print(f"CPF {cpf} encontrado em pelo menos uma linha.")
        linhas_com_cpf = df.loc[df['Nome da Coluna CPF'] == cpf]
        quantidade = len(linhas_com_cpf)
        resultados = linhas_com_cpf.to_dict(orient='records')
        # Exiba todo o conteúdo das linhas onde o CPF desejado aparece
        return  resultados, quantidade
    else:
        print(f"CPF {cpf} não encontrado.")]


# Puxar dados especificos.
def puxar_dados_especifios(cpf, pasta, n_arquivo):
    try:
        arquivo = os.path.join(settings.MEDIA_ROOT, f'planilhas/{pasta}/{n_arquivo}.xlsx')
        df = pd.read_excel(arquivo)
        df['Nome da Coluna CPF'] = df['Nome da Coluna CPF'].astype(str).str.strip()

        # Adiciona zero à esquerda para CPFs com 10 dígitos (ESSA É UMA CORREÇÃO DE ERRO PARA CPF QUE COMEÇÃO COM 0)
        df['Nome da Coluna CPF'] = df['Nome da Coluna CPF'].apply(lambda x: x.zfill(11) if len(x) == 10 else x)

        # Verifica se o CPF fornecido está na planilha
        if (df['Nome da Coluna CPF'] == cpf).any():
            linhas_com_cpf = df.loc[df['Nome da Coluna CPF'] == cpf]
            quantidade = len(linhas_com_cpf)

            # Selecione apenas as colunas desejadas que sejam incluidas na coleta de dados
            colunas_selecionadas = ['Nome da Coluna CPF', 'NOME', 'CNES', 'CBO', 'COMPLEMENTO', 'OBSERVAÇÃO'] # Esses são os titulos das colunas.
            resultados = linhas_com_cpf[colunas_selecionadas].to_dict(orient='records')
            return resultados, quantidade
        else:
            erro = "CPF não encontrado na planilha selecionada! Entre em contato com a TI ou mude o mês e o ano!"
            return erro, 'erro_code'
    except Exception as e:
        erro = e
        return erro, 'erro_code'
    
# Função que retorna para um template especifico, que tem como objetivo servir como pagina de erros.
def error_handler(request, exception=None):
    data = Dados.objects.all()
    
    return render(request, 'error.html', {'erro_code': exception, 'data': data}, status=500)

# Converter DOCX para PDF
def convert_docx_to_pdf(docx_path, pdf_path):
    subprocess.call(['soffice',
                     '--convert-to',
                     'pdf',
                     '--outdir',
                     os.path.dirname(pdf_path),
                     docx_path])

# Abrir PDF direto no navegador pela aplicação DJANGO.
# NO LINUX:
def libreoffice(documentos, pasta_usuario, usuario):
    caminhos_pdf = []

    for doc in documentos:
        docx_path = os.path.join(pasta_usuario, doc)
        pdf_path = docx_path.replace('.docx', '.pdf')
        convert_docx_to_pdf(docx_path, pdf_path)
        caminhos_pdf.append(pdf_path)

    combined_pdf_path = os.path.join(pasta_usuario, f'{usuario}_relatorio.pdf')
    combine_pdfs(caminhos_pdf, combined_pdf_path)

    with open(combined_pdf_path, 'rb') as pdf_file:
        pdf_bytes = pdf_file.read()

    response = HttpResponse(pdf_bytes, content_type='application/pdf')
    response['Content-Disposition'] = f'inline; filename={usuario}_relatorio.pdf'

    for doc in documentos:
        docx_path = os.path.join(pasta_usuario, doc)
        pdf_path = docx_path.replace('.docx', '.pdf')
        os.remove(docx_path)
        os.remove(pdf_path)

    return response

# Juntar PDF's em um unico arquivo. Linux.
def combine_pdfs(pdf_paths, combined_pdf_path):
    subprocess.call(['pdftk'] + pdf_paths + ['cat', 'output', combined_pdf_path])

# No Windows:
def word(documentos, pasta_usuario, usuario):

# Inicializa o COM
    import comtypes
    comtypes.CoInitialize()
    # Converter documentos do Word para PDF
    caminhos_pdf = []
    for doc in documentos:
        caminho_pdf = doc.replace('.docx', '.pdf')
        convert(doc, caminho_pdf)
        caminhos_pdf.append(caminho_pdf)

    # Combinar PDFs
    caminho_pdf_combinado = os.path.join(pasta_usuario, f'{usuario}_relatorio.pdf')
    combinar_pdfs(caminhos_pdf, caminho_pdf_combinado)

    # Ler o conteúdo do arquivo PDF combinado
    with open(caminho_pdf_combinado, 'rb') as pdf_file:
        pdf_bytes = pdf_file.read()

    # Gerar uma resposta HTTP com o conteúdo do PDF
    response = HttpResponse(pdf_bytes, content_type='application/pdf')
    response['Content-Disposition'] = f'inline; filename={usuario}_relatorio.pdf'
    for c in documentos:
        caminho = os.path.join(pasta_usuario, c)
        deletar_arquivo(caminho)

    return response

# Juntar PDF's em um unico arquivo. Windows.
def combinar_pdfs(caminhos_pdfs, caminho_pdf_combinado):
    # Combina os PDFs em um único PDF
    pdf_writer = PdfWriter()
    for caminho_pdf in caminhos_pdfs:
        pdf_reader = PdfReader(caminho_pdf)
        for page in pdf_reader.pages:
            pdf_writer.add_page(page)

    with open(caminho_pdf_combinado, 'wb') as output_pdf:
        pdf_writer.write(output_pdf)

    for caminho_pdf in caminhos_pdfs:
        deletar_arquivo(caminho_pdf)

# Carregar um arquivo pela aplicação para dentro do servidor. UPLOAD
def anexar_planilha(request):
    if request.method == 'POST':
        n_arquivo = request.POST.get('mes')
        pasta = request.POST.get('ano')
        arquivo = request.FILES['arquivo']

        # Define o diretório onde os arquivos serão salvos
        local_do_arquivo = os.path.join(settings.MEDIA_ROOT, f'planilhas/{pasta}/')
        os.makedirs(local_do_arquivo, exist_ok=True)

        # Configura o FileSystemStorage para usar o diretório específico
        fs = FileSystemStorage(location=local_do_arquivo)

        # Configura o nome do arquivo, se necessário
        original_filename = f'{n_arquivo}.xlsx'
        
        # Verifica se o arquivo já existe e o exclui, se necessário
        arquivo_existente = os.path.join(local_do_arquivo, original_filename)
        if os.path.exists(arquivo_existente):
            with open(arquivo_existente, 'wb') as f:
                f.close()
            os.remove(arquivo_existente)
            logger.warning(f"Foi realizada uma substituição, o arquivo {original_filename} existente foi removido.")

        # Salva o novo arquivo
        filename = fs.save(original_filename, arquivo)
        file_url = fs.url(filename)
        logger.warning(f'O usuario {request.user}, adicionou uma nova planilha referente ao mês de {n_arquivo} e ano de {pasta}!')

        # Faça algo com os dados do mês, ano e o arquivo enviado
        return redirect(inicio)
    return render(request, 'anexar_planilha.html')

# Função para anexar uma imagem de assinatura em um documento docx, em até 4 campos.
def adicionar_assinatura(modelo, quantidade, nome_completo, cargo, imagem_assinatura, numero_assinaturas):
    # Calcula o número de linhas e colunas com base na quantidade de assinaturas
    if quantidade == 1:
        paragraph = modelo.add_paragraph()
        run = paragraph.add_run()
        run.add_picture(imagem_assinatura, width=Inches(2.0), height=Inches(1.0))
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        

        # Adicionar uma linha pontilhada após a assinatura
        novo_paragrafo = modelo.add_paragraph('_' * 40)
        novo_paragrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Adicionar o cargo do usuário em um novo parágrafo
        novo_paragrafo = modelo.add_paragraph(nome_completo)
        novo_paragrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Adicionar o cargo do usuário em um novo parágrafo
        novo_paragrafo = modelo.add_paragraph(cargo)
        novo_paragrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Adciona uma linha após a assinatura
        novo_paragrafo = modelo.add_paragraph()

    elif quantidade == 2:
        # Verifica qual é a assinatura
        if numero_assinaturas == 0:
            # Define a tabela
            linhas, colunas = 1, 2  
            # Adiciona uma tabela ao documento
            tabela = modelo.add_table(rows=linhas, cols=colunas)
            linha = 0
            coluna = 0
            paragrafo_assinatura = tabela.cell(linha, coluna).add_paragraph()
            run_assinatura = paragrafo_assinatura.add_run()
            run_assinatura.add_picture(imagem_assinatura, width=Inches(2.0), height=Inches(1.0))
            paragrafo_assinatura.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Adiciona uma linha pontilhada após a assinatura
            tabela.cell(linha, coluna).add_paragraph('_' * 30).alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Adiciona o nome do assinante
            tabela.cell(linha, coluna).add_paragraph(nome_completo).alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Adiciona o cargo do usuário
            tabela.cell(linha, coluna).add_paragraph(cargo).alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Adiciona uma linha em branco
            tabela.cell(linha, coluna).add_paragraph()
        else:
            for tabela in modelo.tables:
                linha = 0
                coluna = 1
                 # Adiciona a assinatura na célula apropriada da tabela
                paragrafo_assinatura = tabela.cell(linha, coluna).add_paragraph()
                run_assinatura = paragrafo_assinatura.add_run()
                run_assinatura.add_picture(imagem_assinatura, width=Inches(2.0), height=Inches(1.0))
                paragrafo_assinatura.alignment = WD_ALIGN_PARAGRAPH.CENTER

                # Adiciona uma linha pontilhada após a assinatura
                tabela.cell(linha, coluna).add_paragraph('_' * 30).alignment = WD_ALIGN_PARAGRAPH.CENTER

                # Adiciona o nome do assinante
                tabela.cell(linha, coluna).add_paragraph(nome_completo).alignment = WD_ALIGN_PARAGRAPH.CENTER

                # Adiciona o cargo do usuário
                tabela.cell(linha, coluna).add_paragraph(cargo).alignment = WD_ALIGN_PARAGRAPH.CENTER

                # Adiciona uma linha em branco
                tabela.cell(linha, coluna).add_paragraph()

    elif quantidade == 3:
        linhas, colunas = 2, 2

                # Verifica qual é a assinatura
        if numero_assinaturas == 0:
            # Define a tabela
            linhas, colunas = 1, 2  
            # Adiciona uma tabela ao documento
            tabela = modelo.add_table(rows=linhas, cols=colunas)
            linha = 0
            coluna = 0
            paragrafo_assinatura = tabela.cell(linha, coluna).add_paragraph()
            run_assinatura = paragrafo_assinatura.add_run()
            run_assinatura.add_picture(imagem_assinatura, width=Inches(2.0), height=Inches(1.0))
            paragrafo_assinatura.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Adiciona uma linha pontilhada após a assinatura
            tabela.cell(linha, coluna).add_paragraph('_' * 30).alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Adiciona o nome do assinante
            tabela.cell(linha, coluna).add_paragraph(nome_completo).alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Adiciona o cargo do usuário
            tabela.cell(linha, coluna).add_paragraph(cargo).alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Adiciona uma linha em branco
            tabela.cell(linha, coluna).add_paragraph()
        elif numero_assinaturas == 1:
            for tabela in modelo.tables:
                linha = 0
                coluna = 1
                 # Adiciona a assinatura na célula apropriada da tabela
                paragrafo_assinatura = tabela.cell(linha, coluna).add_paragraph()
                run_assinatura = paragrafo_assinatura.add_run()
                run_assinatura.add_picture(imagem_assinatura, width=Inches(2.0), height=Inches(1.0))
                paragrafo_assinatura.alignment = WD_ALIGN_PARAGRAPH.CENTER

                # Adiciona uma linha pontilhada após a assinatura
                tabela.cell(linha, coluna).add_paragraph('_' * 30).alignment = WD_ALIGN_PARAGRAPH.CENTER

                # Adiciona o nome do assinante
                tabela.cell(linha, coluna).add_paragraph(nome_completo).alignment = WD_ALIGN_PARAGRAPH.CENTER

                # Adiciona o cargo do usuário
                tabela.cell(linha, coluna).add_paragraph(cargo).alignment = WD_ALIGN_PARAGRAPH.CENTER

                # Adiciona uma linha em branco
                tabela.cell(linha, coluna).add_paragraph()

        elif numero_assinaturas == 2:
            paragraph = modelo.add_paragraph()
            run = paragraph.add_run()
            run.add_picture(imagem_assinatura, width=Inches(2.0), height=Inches(1.0))
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        

            # Adicionar uma linha pontilhada após a assinatura
            novo_paragrafo = modelo.add_paragraph('_' * 40)
            novo_paragrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Adicionar o cargo do usuário em um novo parágrafo
            novo_paragrafo = modelo.add_paragraph(nome_completo)
            novo_paragrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Adicionar o cargo do usuário em um novo parágrafo
            novo_paragrafo = modelo.add_paragraph(cargo)
            novo_paragrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Adciona uma linha após a assinatura
            novo_paragrafo = modelo.add_paragraph()
    
    elif quantidade == 4:
        # Verifica qual é a assinatura
        if numero_assinaturas == 0:
            # Define a tabela
            linhas, colunas = 2, 2  
            # Adiciona uma tabela ao documento
            tabela = modelo.add_table(rows=linhas, cols=colunas)
            linha = 0
            coluna = 0
            paragrafo_assinatura = tabela.cell(linha, coluna).add_paragraph()
            run_assinatura = paragrafo_assinatura.add_run()
            run_assinatura.add_picture(imagem_assinatura, width=Inches(2.0), height=Inches(1.0))
            paragrafo_assinatura.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Adiciona uma linha pontilhada após a assinatura
            tabela.cell(linha, coluna).add_paragraph('_' * 30).alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Adiciona o nome do assinante
            tabela.cell(linha, coluna).add_paragraph(nome_completo).alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Adiciona o cargo do usuário
            tabela.cell(linha, coluna).add_paragraph(cargo).alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Adiciona uma linha em branco
            tabela.cell(linha, coluna).add_paragraph()
        elif numero_assinaturas == 1:
            for tabela in modelo.tables:
                linha = 0
                coluna = 1
                 # Adiciona a assinatura na célula apropriada da tabela
                paragrafo_assinatura = tabela.cell(linha, coluna).add_paragraph()
                run_assinatura = paragrafo_assinatura.add_run()
                run_assinatura.add_picture(imagem_assinatura, width=Inches(2.0), height=Inches(1.0))
                paragrafo_assinatura.alignment = WD_ALIGN_PARAGRAPH.CENTER

                # Adiciona uma linha pontilhada após a assinatura
                tabela.cell(linha, coluna).add_paragraph('_' * 30).alignment = WD_ALIGN_PARAGRAPH.CENTER

                # Adiciona o nome do assinante
                tabela.cell(linha, coluna).add_paragraph(nome_completo).alignment = WD_ALIGN_PARAGRAPH.CENTER

                # Adiciona o cargo do usuário
                tabela.cell(linha, coluna).add_paragraph(cargo).alignment = WD_ALIGN_PARAGRAPH.CENTER

                # Adiciona uma linha em branco
                tabela.cell(linha, coluna).add_paragraph()
        elif numero_assinaturas == 2:
            for tabela in modelo.tables:
                linha = 1
                coluna = 0
                 # Adiciona a assinatura na célula apropriada da tabela
                paragrafo_assinatura = tabela.cell(linha, coluna).add_paragraph()
                run_assinatura = paragrafo_assinatura.add_run()
                run_assinatura.add_picture(imagem_assinatura, width=Inches(2.0), height=Inches(1.0))
                paragrafo_assinatura.alignment = WD_ALIGN_PARAGRAPH.CENTER

                # Adiciona uma linha pontilhada após a assinatura
                tabela.cell(linha, coluna).add_paragraph('_' * 30).alignment = WD_ALIGN_PARAGRAPH.CENTER

                # Adiciona o nome do assinante
                tabela.cell(linha, coluna).add_paragraph(nome_completo).alignment = WD_ALIGN_PARAGRAPH.CENTER

                # Adiciona o cargo do usuário
                tabela.cell(linha, coluna).add_paragraph(cargo).alignment = WD_ALIGN_PARAGRAPH.CENTER

                # Adiciona uma linha em branco
                tabela.cell(linha, coluna).add_paragraph() 
        elif numero_assinaturas == 3:
            for tabela in modelo.tables:
                linha = 1
                coluna = 1
                 # Adiciona a assinatura na célula apropriada da tabela
                paragrafo_assinatura = tabela.cell(linha, coluna).add_paragraph()
                run_assinatura = paragrafo_assinatura.add_run()
                run_assinatura.add_picture(imagem_assinatura, width=Inches(2.0), height=Inches(1.0))
                paragrafo_assinatura.alignment = WD_ALIGN_PARAGRAPH.CENTER

                # Adiciona uma linha pontilhada após a assinatura
                tabela.cell(linha, coluna).add_paragraph('_' * 30).alignment = WD_ALIGN_PARAGRAPH.CENTER

                # Adiciona o nome do assinante
                tabela.cell(linha, coluna).add_paragraph(nome_completo).alignment = WD_ALIGN_PARAGRAPH.CENTER

                # Adiciona o cargo do usuário
                tabela.cell(linha, coluna).add_paragraph(cargo).alignment = WD_ALIGN_PARAGRAPH.CENTER

                # Adiciona uma linha em branco
                tabela.cell(linha, coluna).add_paragraph() 





